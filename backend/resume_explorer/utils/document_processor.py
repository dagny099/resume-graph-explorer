"""
Document Text Extraction Utilities

Extracts text from various document formats (PDF, DOCX, TXT, MD).
Provides fallback mechanisms for robust extraction.
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any
import io

from .logger import logger


class DocumentProcessor:
    """
    Extract text from various document formats.

    Supports:
    - PDF (.pdf) - using PyMuPDF (fitz) with pdfplumber fallback
    - Word (.docx, .doc) - using python-docx
    - Plain text (.txt, .md)
    """

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from document based on file extension.

        Args:
            file_path: Path to document file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension == '.pdf':
            return DocumentProcessor._extract_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return DocumentProcessor._extract_docx(file_path)
        elif extension in ['.txt', '.md', '.markdown']:
            return DocumentProcessor._extract_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
        """
        Extract text from file bytes (for uploaded files).

        Args:
            file_bytes: File content as bytes
            filename: Original filename (for extension detection)

        Returns:
            Extracted text content
        """
        extension = Path(filename).suffix.lower()

        if extension == '.pdf':
            return DocumentProcessor._extract_pdf_from_bytes(file_bytes)
        elif extension in ['.docx', '.doc']:
            return DocumentProcessor._extract_docx_from_bytes(file_bytes)
        elif extension in ['.txt', '.md', '.markdown']:
            return file_bytes.decode('utf-8', errors='ignore')
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        """
        Extract text from PDF using PyMuPDF with pdfplumber fallback.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text
        """
        try:
            import fitz  # PyMuPDF

            text_parts = []
            with fitz.open(file_path) as doc:
                for page_num, page in enumerate(doc, start=1):
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(text)

            result = '\n\n'.join(text_parts)

            if not result.strip():
                logger.warning(f"PyMuPDF extracted empty text from {file_path}, trying pdfplumber")
                return DocumentProcessor._extract_pdf_fallback(file_path)

            logger.info(f"Extracted {len(result)} characters from PDF using PyMuPDF")
            return result

        except ImportError:
            logger.warning("PyMuPDF not available, using pdfplumber")
            return DocumentProcessor._extract_pdf_fallback(file_path)
        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {e}, trying pdfplumber")
            return DocumentProcessor._extract_pdf_fallback(file_path)

    @staticmethod
    def _extract_pdf_fallback(file_path: str) -> str:
        """
        Fallback PDF extraction using pdfplumber.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text
        """
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            result = '\n\n'.join(text_parts)
            logger.info(f"Extracted {len(result)} characters from PDF using pdfplumber")
            return result

        except ImportError:
            raise ImportError("Neither PyMuPDF nor pdfplumber is installed. "
                            "Install with: pip install PyMuPDF pdfplumber")
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}")
            raise

    @staticmethod
    def _extract_pdf_from_bytes(file_bytes: bytes) -> str:
        """
        Extract text from PDF bytes.

        Args:
            file_bytes: PDF file content

        Returns:
            Extracted text
        """
        try:
            import fitz  # PyMuPDF

            text_parts = []
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page in doc:
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(text)

            result = '\n\n'.join(text_parts)

            if not result.strip():
                logger.warning("PyMuPDF extracted empty text from bytes, trying pdfplumber")
                return DocumentProcessor._extract_pdf_from_bytes_fallback(file_bytes)

            return result

        except ImportError:
            logger.warning("PyMuPDF not available, using pdfplumber")
            return DocumentProcessor._extract_pdf_from_bytes_fallback(file_bytes)
        except Exception as e:
            logger.error(f"PyMuPDF extraction from bytes failed: {e}")
            return DocumentProcessor._extract_pdf_from_bytes_fallback(file_bytes)

    @staticmethod
    def _extract_pdf_from_bytes_fallback(file_bytes: bytes) -> str:
        """
        Fallback PDF extraction from bytes using pdfplumber.

        Args:
            file_bytes: PDF file content

        Returns:
            Extracted text
        """
        try:
            import pdfplumber

            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

            return '\n\n'.join(text_parts)

        except ImportError:
            raise ImportError("Neither PyMuPDF nor pdfplumber is installed")
        except Exception as e:
            logger.error(f"pdfplumber extraction from bytes failed: {e}")
            raise

    @staticmethod
    def _extract_docx(file_path: str) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text
        """
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            result = '\n\n'.join(text_parts)
            logger.info(f"Extracted {len(result)} characters from DOCX")
            return result

        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    @staticmethod
    def _extract_docx_from_bytes(file_bytes: bytes) -> str:
        """
        Extract text from DOCX bytes.

        Args:
            file_bytes: DOCX file content

        Returns:
            Extracted text
        """
        try:
            from docx import Document

            doc = Document(io.BytesIO(file_bytes))
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            return '\n\n'.join(text_parts)

        except ImportError:
            raise ImportError("python-docx not installed")
        except Exception as e:
            logger.error(f"DOCX extraction from bytes failed: {e}")
            raise

    @staticmethod
    def _extract_text_file(file_path: str) -> str:
        """
        Extract text from plain text file.

        Args:
            file_path: Path to text file

        Returns:
            File content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            logger.info(f"Read {len(content)} characters from text file")
            return content

        except UnicodeDecodeError:
            # Try with latin-1 encoding as fallback
            logger.warning("UTF-8 decoding failed, trying latin-1")
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            raise

    @staticmethod
    def get_document_metadata(file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from document.

        Args:
            file_path: Path to document

        Returns:
            Dictionary with metadata (filename, size, extension, etc.)
        """
        path = Path(file_path)

        metadata = {
            'filename': path.name,
            'extension': path.suffix.lower(),
            'size_bytes': path.stat().st_size,
            'size_kb': round(path.stat().st_size / 1024, 2),
        }

        # Try to get page count for PDFs
        if path.suffix.lower() == '.pdf':
            try:
                import fitz
                with fitz.open(file_path) as doc:
                    metadata['page_count'] = len(doc)
            except Exception:
                pass

        return metadata


__all__ = ['DocumentProcessor']
